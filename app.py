import streamlit as st
import pandas as pd
from docx import Document
import io
from datetime import datetime
import hashlib

# Page configuration
st.set_page_config(
    page_title="Bonafide Certificate Generator",
    page_icon="📜",
    layout="centered"
)

# Password authentication
def check_password():
    """Returns True if the user has entered the correct password."""

    # Get password from Streamlit secrets, or use default for local testing
    try:
        correct_password = st.secrets.get("password", "school2024")
    except:
        correct_password = "school2024"  # Default password for local testing

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if hashlib.sha256(st.session_state["password"].encode()).hexdigest() == hashlib.sha256(correct_password.encode()).hexdigest():
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store password
        else:
            st.session_state["password_correct"] = False

    # First run or password not yet entered
    if "password_correct" not in st.session_state:
        st.markdown("### 🔐 Authentication Required")
        st.markdown("Please enter the password to access the certificate generator.")
        st.text_input(
            "Password",
            type="password",
            on_change=password_entered,
            key="password",
            help="Contact school administration for the password"
        )
        st.info("💡 Default password for testing: `school2024`")
        return False

    # Password incorrect
    elif not st.session_state["password_correct"]:
        st.markdown("### 🔐 Authentication Required")
        st.markdown("Please enter the password to access the certificate generator.")
        st.text_input(
            "Password",
            type="password",
            on_change=password_entered,
            key="password",
            help="Contact school administration for the password"
        )
        st.error("❌ Incorrect password. Please try again.")
        return False

    # Password correct
    else:
        return True

# Check authentication before showing the app
if not check_password():
    st.stop()

# Custom CSS for white-label commercial app (remove ALL branding)
st.markdown("""
    <style>
    /* === COMPLETE BRANDING REMOVAL === */

    /* Hide ALL Streamlit UI elements */
    #MainMenu {visibility: hidden !important;}
    footer {visibility: hidden !important;}
    header {visibility: hidden !important;}
    .stDeployButton {display: none !important;}
    .stDecoration {display: none !important;}

    /* Hide GitHub/Fork/Deploy buttons */
    button[title="View app source on GitHub"] {display: none !important;}
    button[data-testid="baseButton-header"] {display: none !important;}
    a[href*="github.com"] {display: none !important;}
    [data-testid="stToolbar"] {display: none !important;}
    [data-testid="stDecoration"] {display: none !important;}
    [data-testid="stStatusWidget"] {display: none !important;}

    /* Hide footer with Streamlit branding */
    footer {display: none !important;}
    footer > div {display: none !important;}
    .viewerBadge_container__1QSob {display: none !important;}
    .viewerBadge_link__1S137 {display: none !important;}
    .viewerBadge_text__1JaDK {display: none !important;}
    div[class*="viewerBadge"] {display: none !important;}

    /* Hide header and all header buttons */
    [data-testid="stHeader"] {display: none !important;}
    .stApp > header {display: none !important;}
    header[data-testid="stHeader"] {display: none !important;}

    /* Hide profile/avatar elements */
    img[alt="App Creator Avatar"] {display: none !important;}
    img[class*="profileImage"] {display: none !important;}
    img[class*="appCreatorAvatar"] {display: none !important;}
    img[src*="avatar"] {display: none !important;}
    div[class*="profileContainer"] {display: none !important;}
    div[class*="profilePreview"] {display: none !important;}
    div[class*="_profileContainer_"] {display: none !important;}
    div[class*="_profileImage_"] {display: none !important;}
    a[href*="user/"] {display: none !important;}
    a[href*="/user/"] {display: none !important;}
    a[target="_blank"][rel*="noopener"] {display: none !important;}

    /* Hide ALL elements with branding-related keywords */
    [class*="profile"] {display: none !important;}
    [class*="avatar"] {display: none !important;}
    [class*="Avatar"] {display: none !important;}
    [class*="Profile"] {display: none !important;}
    [class*="viewerBadge"] {display: none !important;}
    [class*="ViewerBadge"] {display: none !important;}

    /* Hide Streamlit status/analytics */
    script[src*="streamlitstatus.com"] {display: none !important;}
    script[src*="segment"] {display: none !important;}
    iframe[class*="statusFrame"] {display: none !important;}
    iframe[src*="streamlit"] {display: none !important;}

    /* Hide bottom-right corner elements */
    .styles_viewerBadge__1yB5_ {display: none !important;}
    [data-testid="collapsedControl"] {display: none !important;}
    .stActionButton {display: none !important;}
    button[kind="header"] {display: none !important;}
    button[kind="headerNoPadding"] {display: none !important;}

    /* Hide iframes and external links */
    .main > div > div > div > iframe {display: none !important;}
    iframe[title*="Streamlit"] {display: none !important;}

    /* Hide dynamic class elements (gzau, etc.) */
    div[class*="gzau"] > div {display: none !important;}
    div[class*="gzau"] > a {display: none !important;}
    div[class*="gzau"] > img {display: none !important;}
    a[class*="gzau"] {display: none !important;}

    /* Hide "Made with Streamlit" and similar */
    a[title*="Streamlit"] {display: none !important;}
    a[aria-label*="Streamlit"] {display: none !important;}
    div[aria-label*="Streamlit"] {display: none !important;}

    /* Remove top padding where header was */
    .block-container {
        padding-top: 1rem !important;
    }

    /* App styling */
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
        padding: 0.5rem;
        font-size: 1.1rem;
    }
    .upload-section {
        background-color: #f0f2f6;
        padding: 1.5rem;
        border-radius: 0.5rem;
        margin-bottom: 2rem;
    }
    .success-message {
        padding: 1rem;
        background-color: #d4edda;
        border-radius: 0.5rem;
        color: #155724;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Title and description
st.title("📜 Bonafide Certificate Generator")
st.markdown("---")

# Initialize session state for files
if 'excel_file' not in st.session_state:
    st.session_state.excel_file = None
if 'word_file' not in st.session_state:
    st.session_state.word_file = None
if 'df' not in st.session_state:
    st.session_state.df = None

# File Upload Section
st.markdown("## 📁 Step 1: Upload Required Files")
st.markdown('<div class="upload-section">', unsafe_allow_html=True)

col1, col2 = st.columns(2)

with col1:
    st.markdown("#### 📊 Student Records (Excel)")
    uploaded_excel = st.file_uploader(
        "Upload Excel File",
        type=['xlsx', 'xls'],
        help="Upload the student records Excel file",
        key="excel_uploader"
    )

    if uploaded_excel:
        st.session_state.excel_file = uploaded_excel
        try:
            # Load Excel file
            st.session_state.df = pd.read_excel(uploaded_excel, header=1)
            st.success(f"✅ Loaded {len(st.session_state.df)} student records")
        except Exception as e:
            st.error(f"❌ Error reading Excel file: {str(e)}")
            st.session_state.df = None

with col2:
    st.markdown("#### 📄 Certificate Template (Word)")
    uploaded_word = st.file_uploader(
        "Upload Word Template",
        type=['docx'],
        help="Upload the bonafide certificate template",
        key="word_uploader"
    )

    if uploaded_word:
        st.session_state.word_file = uploaded_word
        st.success("✅ Template loaded successfully")

st.markdown('</div>', unsafe_allow_html=True)

# Check if files are uploaded
if st.session_state.excel_file is None or st.session_state.word_file is None:
    st.warning("⚠️ Please upload both files to continue")

    # Show instructions
    with st.expander("📖 How to use this app"):
        st.markdown("""
        ### Instructions:

        1. **Upload Student Records** (Excel file)
           - Must contain columns: Reg. No, Name, Father, Class, Session, DOB, Gender
           - First row should be school name, second row should be column headers

        2. **Upload Certificate Template** (Word file)
           - Must contain placeholders: {{REGNO}}, {{NAME}}, {{FATHER}}, {{CLASS}}, etc.

        3. **Enter Registration Number** and generate certificate

        4. **Download** the generated certificate

        ### Note:
        - Files are stored only during your session
        - When you close the browser, files are deleted
        - You'll need to upload again next time you use the app
        """)

    st.stop()

# Show statistics if data is loaded
if st.session_state.df is not None:
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Students", len(st.session_state.df))
    with col2:
        st.metric("Total Classes", st.session_state.df["Class"].nunique())
    with col3:
        st.metric("Sessions", st.session_state.df["Session"].nunique())

# Generate Certificate Section
st.markdown("---")
st.markdown("## 🔍 Step 2: Generate Certificate")

# Input section
col1, col2 = st.columns([2, 1])

with col1:
    reg_input = st.text_input(
        "Enter Registration Number",
        placeholder="e.g., 2552",
        help="Enter the student's registration number (Reg. No)",
        key="reg_input"
    )

with col2:
    st.markdown("<br>", unsafe_allow_html=True)
    search_button = st.button("🔍 Search & Generate", use_container_width=True)

# Generate certificate function
def generate_certificate(row, reg_no, word_template):
    try:
        # Extract student information
        name = row["Name"]
        father = row["Father"]
        class_ = row["Class"]
        session = row["Session"]
        dob = row["DOB"]
        gender = str(row["Gender"]).strip().upper()

        # Gender logic
        if gender == "M":
            relation = "S/O"
            pronoun_subject = "He"
            pronoun_possessive = "His"
        else:
            relation = "D/O"
            pronoun_subject = "She"
            pronoun_possessive = "Her"

        # Load Word template from uploaded file
        doc = Document(word_template)

        # Create replacement dictionary
        replacements = {
            "{{NAME}}": name,
            "{{RELATION}}": relation,
            "{{FATHER}}": father,
            "{{CLASS}}": str(class_),
            "{{SESSION}}": session,
            "{{DOB}}": str(dob),
            "{{PRONOUN_SUBJECT}}": pronoun_subject,
            "{{PRONOUN_POSSESSIVE}}": pronoun_possessive
        }

        def replace_in_paragraph(para):
            """Replace placeholders in paragraph, handling split runs"""
            for key, value in replacements.items():
                if key in para.text:
                    # Merge all run texts
                    full_text = para.text

                    # Replace all placeholders
                    new_text = full_text
                    for k, v in replacements.items():
                        new_text = new_text.replace(k, v)

                    # If text changed, update the paragraph
                    if new_text != full_text:
                        # Clear existing runs and add new text
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = new_text
                        else:
                            para.add_run(new_text)
                    break

        # Replace in all paragraphs
        for para in doc.paragraphs:
            replace_in_paragraph(para)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)

        # Replace in headers and footers
        for section in doc.sections:
            # Replace in header
            for para in section.header.paragraphs:
                replace_in_paragraph(para)

            # Replace in footer
            for para in section.footer.paragraphs:
                replace_in_paragraph(para)

        # Save to BytesIO object
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io, name, None
    except Exception as e:
        return None, None, f"Error generating certificate: {str(e)}"

# Search and generate certificate
if search_button:
    if not reg_input:
        st.warning("⚠️ Please enter a registration number.")
    else:
        reg_input = reg_input.strip()

        with st.spinner("🔄 Searching for student..."):
            # Filter student record
            df = st.session_state.df
            student = df[df["Reg. No"].fillna(0).astype(int).astype(str) == reg_input]

            if student.empty:
                st.error(f"❌ No student found with Registration Number: {reg_input}")
                st.info("💡 Tip: Make sure you're entering the Reg. No, not the St. ID")
            else:
                # Get first (and only) row
                row = student.iloc[0]

                # Display student information
                st.success(f"✅ Student Found!")

                # Student details card
                st.markdown("### Student Details")
                info_col1, info_col2 = st.columns(2)

                with info_col1:
                    st.write(f"**Name:** {row['Name']}")
                    st.write(f"**Father's Name:** {row['Father']}")
                    st.write(f"**Class:** {row['Class']}")

                with info_col2:
                    st.write(f"**Reg. No:** {reg_input}")
                    st.write(f"**Session:** {row['Session']}")
                    st.write(f"**DOB:** {row['DOB']}")

                st.markdown("---")

                # Generate certificate
                with st.spinner("📝 Generating certificate..."):
                    doc_io, name, error = generate_certificate(row, reg_input, st.session_state.word_file)

                    if error:
                        st.error(error)
                    else:
                        st.success("🎉 Certificate generated successfully!")

                        # Download button
                        st.download_button(
                            label="📥 Download Certificate",
                            data=doc_io,
                            file_name=f"TO WHOM IT MAY CONCERN - {reg_input}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

# Footer
st.markdown("---")
st.markdown("""
    <div style='text-align: center; color: #666; padding: 1rem;'>
        <small>Certificate Generator v2.0<br>
        🔒 Secure • Private • Professional</small>
    </div>
""", unsafe_allow_html=True)
